from datetime import datetime, timezone
from pathlib import Path

import docx

from spec_validator.models.spec import (
    DataType,
    FieldSpec,
    NamingConvention,
    NullPolicy,
    SpecDocument,
)
from spec_validator.parsers.base import BaseSpecParser

_FIELD_NAME_HEADERS = {"field", "name", "column", "field_name", "column_name", "fieldname"}
_TYPE_HEADERS = {"type", "data_type", "datatype", "dtype"}
_NULLABLE_HEADERS = {"nullable", "required", "mandatory", "null", "optional"}
_ALLOWED_HEADERS = {"allowed_values", "values", "enum", "valid_values", "allowed"}
_DESCRIPTION_HEADERS = {"description", "desc", "notes", "note", "comment"}
_PATTERN_HEADERS = {"pattern", "regex", "format"}

_TYPE_MAP: dict[str, DataType] = {
    "string": DataType.STRING, "str": DataType.STRING, "text": DataType.STRING,
    "integer": DataType.INTEGER, "int": DataType.INTEGER,
    "float": DataType.FLOAT, "double": DataType.FLOAT, "decimal": DataType.FLOAT,
    "boolean": DataType.BOOLEAN, "bool": DataType.BOOLEAN,
    "date": DataType.DATE,
    "datetime": DataType.DATETIME, "timestamp": DataType.DATETIME,
    "email": DataType.EMAIL,
    "url": DataType.URL, "uri": DataType.URL,
    "enum": DataType.ENUM,
}


class WordSpecParser(BaseSpecParser):
    def can_parse(self, path: str) -> bool:
        return Path(path).suffix.lower() == ".docx"

    def parse(self, path: str, spec_id: str) -> SpecDocument:
        document = docx.Document(path)
        fields = self._tables_to_field_specs(document)
        raw_content = self._serialize_to_text(document)
        naming_convention = self._extract_naming_convention(document)
        title = self._extract_title(document)

        return SpecDocument(
            spec_id=spec_id,
            source_path=str(Path(path).resolve()),
            source_format="docx",
            title=title or Path(path).stem,
            fields=fields,
            naming_convention=naming_convention,
            raw_content=raw_content,
            parsed_at=datetime.now(timezone.utc).isoformat(),
        )

    def _extract_title(self, document) -> str:
        for para in document.paragraphs:
            if para.style.name.startswith("Heading 1") and para.text.strip():
                return para.text.strip()
        return ""

    def _tables_to_field_specs(self, document) -> list[FieldSpec]:
        for table in document.tables:
            if not table.rows:
                continue
            header_row = [cell.text.strip() for cell in table.rows[0].cells]
            if not any(h.lower() in _FIELD_NAME_HEADERS for h in header_row):
                continue

            rows = []
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(dict(zip(header_row, cells)))

            fields = self._rows_to_field_specs(header_row, rows)
            if fields:
                return fields
        return []

    def _rows_to_field_specs(self, headers: list[str], rows: list[dict]) -> list[FieldSpec]:
        def _find_key(candidates: set[str]) -> str | None:
            for h in headers:
                normalized = h.lower().strip().replace(" ", "_")
                if normalized in candidates:
                    return h
            return None

        name_key = _find_key(_FIELD_NAME_HEADERS)
        type_key = _find_key(_TYPE_HEADERS)
        nullable_key = _find_key(_NULLABLE_HEADERS)
        allowed_key = _find_key(_ALLOWED_HEADERS)
        desc_key = _find_key(_DESCRIPTION_HEADERS)
        pattern_key = _find_key(_PATTERN_HEADERS)

        if name_key is None:
            return []

        fields = []
        for row in rows:
            name = row.get(name_key, "").strip()
            if not name:
                continue

            raw_type = row.get(type_key, "").lower().strip() if type_key else ""
            data_type = _TYPE_MAP.get(raw_type, DataType.ANY)

            raw_nullable = row.get(nullable_key, "").lower().strip() if nullable_key else ""
            if raw_nullable in ("yes", "true", "1", "nullable", "null"):
                nullable = NullPolicy.NULLABLE
            elif raw_nullable in ("no", "false", "0", "required", "mandatory", "not null"):
                nullable = NullPolicy.REQUIRED
            else:
                nullable = NullPolicy.OPTIONAL

            raw_allowed = row.get(allowed_key, "").strip() if allowed_key else ""
            allowed_values = None
            if raw_allowed:
                allowed_values = [v.strip() for v in raw_allowed.split(",") if v.strip()]

            description = row.get(desc_key, "").strip() if desc_key else ""
            pattern = row.get(pattern_key, "").strip() if pattern_key else ""

            fields.append(FieldSpec(
                name=name,
                data_type=data_type,
                nullable=nullable,
                allowed_values=allowed_values or None,
                pattern=pattern or None,
                description=description,
            ))
        return fields

    def _serialize_to_text(self, document) -> str:
        parts = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for i, table in enumerate(document.tables):
            parts.append(f"\n[Table {i + 1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("| " + " | ".join(cells) + " |")

        return "\n".join(parts)

    def _extract_naming_convention(self, document) -> NamingConvention | None:
        in_naming_section = False
        description_parts: list[str] = []

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                in_naming_section = "naming" in text.lower() or "convention" in text.lower()
            elif in_naming_section:
                description_parts.append(text)

        if not description_parts:
            return None

        desc = " ".join(description_parts)
        pattern = None
        if "snake_case" in desc.lower():
            pattern = "snake_case"
        elif "camelcase" in desc.lower() or "camel_case" in desc.lower():
            pattern = "camelCase"
        elif "screaming_snake" in desc.lower():
            pattern = "SCREAMING_SNAKE_CASE"
        elif "pascalcase" in desc.lower() or "pascal_case" in desc.lower():
            pattern = "PascalCase"

        return NamingConvention(pattern=pattern, description=desc)
